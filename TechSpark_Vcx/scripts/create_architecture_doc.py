from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "phase2_architecture_and_ai_design.docx"
BLUE, NAVY, TEAL, PALE, RED = "2E74B5", "0B2545", "167D71", "E8EEF5", "9B1C1C"

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tc_pr.append(shd)

def margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc.get_or_add_tcPr(); m = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{side}"); node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa"); m.append(node)
    tc.append(m)

def set_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr(); w = tc_pr.find(qn("w:tcW"))
    if w is None:
        w = OxmlElement("w:tcW")
        tc_pr.append(w)
    w.set(qn("w:w"), str(width)); w.set(qn("w:type"), "dxa")

def style_table(table, widths):
    table.autofit = False
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_width(cell, widths[i]); margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.1
                for run in p.runs: run.font.name = "Calibri"; run.font.size = Pt(9.5)

def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"; style_table(table, widths)
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value; shade(cell, PALE)
        for run in cell.paragraphs[0].runs: run.bold = True; run.font.color.rgb = RGBColor.from_string(NAVY)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row): cell.text = value
    doc.add_paragraph()

def add_flow(doc, title, nodes):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(nodes)); table.style = "Table Grid"; style_table(table, [9360 // len(nodes)] * len(nodes))
    for i, (label, detail, color) in enumerate(nodes):
        cell = table.cell(0, i); shade(cell, color); p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(label + "\n").bold = True; p.add_run(detail)
        if i < len(nodes) - 1: p.add_run("  →")
    doc.add_paragraph("Figure: " + title + ". All runtime components operate on the local device after setup.", style="Caption")

doc = Document()
sec = doc.sections[0]; sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1); sec.header_distance = sec.footer_distance = Inches(.492)
styles = doc.styles
styles["Normal"].font.name = "Calibri"; styles["Normal"].font.size = Pt(11); styles["Normal"].paragraph_format.space_after = Pt(6); styles["Normal"].paragraph_format.line_spacing = 1.1
for name, size, color, before, after in [("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",12,"1F4D78",8,4)]:
    s=styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
header = sec.header.paragraphs[0]; header.text = "NYAYA SAHAYAK  |  PHASE 2 ARCHITECTURE & AI DESIGN"; header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in header.runs: r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string("6B7785")
footer = sec.footer.paragraphs[0]; footer.text = "Hackathon MVP | Local legal-information guidance only"; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in footer.runs: r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string("6B7785")

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(20); p.paragraph_format.space_after=Pt(4); r=p.add_run("NYAYA SAHAYAK"); r.bold=True; r.font.size=Pt(25); r.font.color.rgb=RGBColor.from_string(NAVY)
p=doc.add_paragraph("Architecture & AI Design | Day 1, Phase 2"); p.runs[0].font.size=Pt(14); p.runs[0].font.color.rgb=RGBColor.from_string("506174")
doc.add_paragraph("A local, offline-first legal guidance MVP for consumer complaints, criminal procedure basics, and family-law basics.")
add_table(doc,["Decision","Selected approach"],[
    ("Runtime","React + FastAPI on localhost; no hosted LLM/API calls"),("Generation","Qwen3-4B-Instruct-2507; Q4 GGUF at runtime; optional QLoRA adapter"),("Retrieval","ChromaDB with sentence-transformers/all-MiniLM-L6-v2 embeddings"),("Privacy","Anonymous, in-memory session only; no persistent chat history"),("Safety","Risk routing before retrieval or generation; citations restricted to retrieved records")],[2700,6660])

doc.add_heading("1. Product Objective and Guardrails", level=1)
doc.add_paragraph("The system gives plain-language, source-grounded legal information. It is not a substitute for a licensed legal professional. The MVP must refuse unsupported claims, avoid case-specific strategy, and direct emergencies to human support.")
add_table(doc,["Supported","Out of scope / escalation"],[
    ("Consumer complaints; FIR and bail background; marriage registration and dowry-law basics.","Immediate danger, domestic violence now, arrest in progress, active court-case strategy, and unrelated legal topics.")],[4680,4680])

add_flow(doc,"High-Level Local Architecture",[("Citizen browser","React + TypeScript", "E8EEF5"),("Local API","FastAPI safety + routing", "DFF2EE"),("RAG services","MiniLM + ChromaDB", "E8EEF5"),("Local SLM","Qwen3 4B Q4", "DFF2EE"),("Rendered answer","sources + confidence", "E8EEF5")])

doc.add_heading("2. Offline Deployment and Data Flow", level=1)
doc.add_paragraph("Setup may download model weights and public statutory materials. After weights and corpus are present locally, the web UI communicates only with the FastAPI service over localhost. The backend reads the local vector database and local GGUF model; user questions are never transmitted to an external service.")
add_flow(doc,"Corpus Ingestion Pipeline",[("Official text","India Code / official source", "E8EEF5"),("Validate","metadata + section boundary", "DFF2EE"),("Chunk","section-aligned passages", "E8EEF5"),("Embed","all-MiniLM-L6-v2", "DFF2EE"),("Persist","Chroma local index", "E8EEF5")])

doc.add_heading("3. Runtime AI and Safety Pipeline", level=1)
add_flow(doc,"Query Processing Pipeline",[("Text query","anonymous session", "E8EEF5"),("Safety gate","crisis / active-case check", "FBE9E9"),("Classify","consumer / criminal / family", "DFF2EE"),("Retrieve","domain-filtered top-k", "E8EEF5"),("Generate + validate","local Qwen; citation check", "DFF2EE"),("Deliver","answer / steps / confidence", "E8EEF5")])
doc.add_paragraph("Safety is deterministic and precedes all model work. An immediate-risk or active-case trigger returns escalation guidance directly. For ordinary queries, the classifier selects one of three domains; only that domain's passages are retrieved. The generator is instructed to use only supplied context. Post-processing rejects citations absent from retrieved metadata. Low retrieval confidence returns a professional-consultation message instead of an answer.")

doc.add_heading("4. Retrieval, Generation, and Confidence Controls", level=1)
add_table(doc,["Control","Implementation rule"],[
    ("Grounding","Use only top-k retrieved records passed to the prompt; never invent an Act, section, or case citation."),("Citation validation","Each displayed citation must map to one retrieved record ID, Act name, section number, and source URL."),("Confidence","Use the best retrieval score; below 0.40, return low-confidence escalation rather than legal guidance."),("Safe first run","If the local Qwen GGUF is not yet present, return a clearly labelled retrieval-safe fallback; never simulate model output."),("Follow-up","Client sends the same transient session ID; no history is written to disk.")],[2700,6660])

doc.add_heading("5. Data and Vector Store Design", level=1)
add_table(doc,["Entity","Fields","Purpose"],[
    ("SourceRecord","id, domain, act_name, section_number, section_text, jurisdiction, source_url","Auditable statutory source unit."),("Chroma document","embedding_id, chunk_text, domain, act_name, section_number, source_url","Local semantic retrieval with domain filtering."),("Instruction example","instruction, context, response, cited_sections","Future QLoRA training data, validated against corpus."),("Transient session","browser-held session ID and current turn context","Follow-up UX without persistent personal data.")],[1850,4300,3210])

doc.add_heading("6. API Contract", level=1)
add_table(doc,["Endpoint","Method","Responsibility","Response"],[
    ("/api/v1/health","GET","Local readiness check","Model/index availability and corpus count."),("/api/v1/chat","POST","Safety route, classify, retrieve, generate/validate","Answer, steps, citations, confidence, disclaimer."),("/api/v1/corpus/reindex","POST","Maintenance-only local rebuild","Indexed record count."),("/api/v1/sources/{id}","GET","Resolve a citation shown in UI","Full local source record.")],[2100,1050,3400,2810])

doc.add_heading("7. Technology and Model Selection", level=1)
add_table(doc,["Layer","Choice","Why it fits the MVP"],[
    ("Frontend","React, TypeScript, Vite","Fast citizen-facing chat interface with strong local development ergonomics."),("Backend","FastAPI, Pydantic","Typed local REST API and testable safety/retrieval modules."),("Embeddings","sentence-transformers/all-MiniLM-L6-v2","Small, fast English embedding model suitable for local search."),("Generator","Qwen/Qwen3-4B-Instruct-2507","Non-gated Apache-2.0 4B instruct model; viable for QLoRA and Q4 local inference."),("Vector store","ChromaDB","Persistent local metadata and vector search."),("Fine-tuning","QLoRA + PEFT, optional","Adapter training after RAG baseline passes tests; avoids risking the core demo.")],[1600,2900,4860])

doc.add_heading("8. Dataset Selection and Provenance", level=1)
doc.add_paragraph("The curated MVP corpus starts with official statutory material: Consumer Protection Act, 2019; BNSS, 2023; Dowry Prohibition Act, 1961; relevant national marriage-law material. Every record retains its source URL. State-dependent requirements are labelled as such and prompt users to confirm the state rather than presenting a universal procedure.")
doc.add_paragraph("Fine-tuning data is not treated as authority. Each instruction/context/response example must cite only statutes that appear verbatim in its context. The training adapter is accepted only after source-citation and unsafe-query evaluation passes.")

doc.add_heading("9. Failure Modes and Demo Acceptance", level=1)
add_table(doc,["Scenario","Expected safe result"],[
    ("No local model file","Health reports degraded; RAG-safe fallback remains labelled and cited."),("Low retrieval relevance","No substantive answer; recommend qualified legal support."),("Crisis / immediate danger","Bypass retrieval and generation; show urgent escalation guidance."),("Unsupported citation","Validation removes/rejects it; response is not presented as sourced."),("No internet after setup","Chat, embedding index, and local model operate only from local assets.")],[3300,6060])
doc.add_paragraph("Phase 3 verification requires API tests for all three domains, out-of-scope requests, low confidence, and crisis routing; a browser end-to-end check; offline operation after setup; and visual review of this document.")

OUT.parent.mkdir(parents=True, exist_ok=True); doc.save(OUT); print(OUT)
